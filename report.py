import customtkinter
import tkinter
import re
from docxtpl import DocxTemplate
from CTkMessagebox import CTkMessagebox
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from hPyT import *
from PIL import Image
import os
from tkcalendar import Calendar
import datetime
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from CTkToolTip import *
from docx import *

import CTkSelectDate
from CTkSelectDate import *
from dataBase import DatabaseManager
import tkinter as tk
from tkinter import filedialog
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")




class Reports(customtkinter.CTkToplevel):
    """
        Класс для создания окна формирования отчетов.

        Attributes:
            frame (customtkinter.CTkScrollableFrame): Прокручиваемый фрейм для размещения виджетов.
            first (FirstFrameChoise): Экземпляр фрейма для выбора параметров отчетов.
            second (FirstFrameReport): Экземпляр фрейма для отчета по статусам работы.
            third (SecondFrameReport): Экземпляр фрейма для отчета по ремонтам за промежуток дат.
            fourth (ThirdFrameReport): Экземпляр фрейма для отчета на выбранное устройство.
            fifth (FourthFrameReport): Экземпляр фрейма для отчета по месту установки.
            sixth (FifthFrameReport): Экземпляр фрейма для отчета по всем устройствам филиала.
            seventh (SixthFrameReport): Экземпляр фрейма для отчета по всем устройствам по структурному подразделению.

        Methods:
            create_widgets: Создает и размещает виджеты в окне.
        """

    def __init__(self, *args,db_manager=None, **kwargs):
        """
                Инициализирует объект Reports.

                Args:
                    *args: Дополнительные аргументы.
                    **kwargs: Дополнительные именованные аргументы.
                """
        super().__init__(*args, **kwargs)

        self.db_manager = db_manager
        self.create_widgets()

    def create_widgets(self):
        """
                Создает и размещает виджеты в окне формирования отчетов.
                """
        self.title("Формирование отчетов")
        self.geometry("1190x650")
        self.minsize(1190, 650)
        self.frame = customtkinter.CTkScrollableFrame(self, height=490)

        self.frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.frame.grid_columnconfigure((0, 1), weight=1)

        self.first = FirstFrameChoise(self.frame,db_manager=self.db_manager)
        self.first.grid(row=0, column=0, padx=10, pady=10, sticky="", columnspan=3)

        # Отчет по статусам работы  - готово
        self.second = FirstFrameReport(self.frame, self.first,db_manager=self.db_manager)
        self.second.grid(row=1, column=0, padx=10, pady=10, sticky="")
        # Отчет по ремонтам за промежуток дат  - готово
        self.third = SecondFrameReport(self.frame, self.first, self.second,db_manager=self.db_manager)
        self.third.grid(row=1, column=1, padx=10, pady=10, sticky="")
        # Отчет на выбранное устройство  - готово
        self.fourth = ThirdFrameReport(self.frame, self.first, self.second,db_manager=self.db_manager)
        self.fourth.grid(row=1, column=2, padx=10, pady=10, sticky="ns")

        # Отчет по месту установки  - готово
        self.fifth = FourthFrameReport(self.frame, self.first,db_manager=self.db_manager)
        self.fifth.grid(row=2, column=0, padx=10, pady=10, sticky="ns")
        # По всем устройствам филиала
        self.sixth = FifthFrameReport(self.frame, self.first,db_manager=self.db_manager)
        self.sixth.grid(row=2, column=1, padx=10, pady=10, sticky="ns")
        # отчет по всем устройствам по структурного подразделения
        self.seventh = SixthFrameReport(self.frame, self.first,db_manager=self.db_manager)
        self.seventh.grid(row=2, column=2, padx=10, pady=10, sticky="ns")


class FirstFrameChoise(customtkinter.CTkFrame):
    """
            Класс для создания фрейма выбора филиала и структурного подразделения.

            Attributes:
                combobox1_branch_office (customtkinter.CTkComboBox): Комбобокс для выбора филиала.
                combobox2_structural_unit (customtkinter.CTkComboBox): Комбобокс для выбора структурного подразделения.

            Methods:
                load_data: Загружает данные для структурного подразделения на основе выбранного филиала.
                FillComboBox: Заполняет комбобокс данными.
            """

    def __init__(self, *args,db_manager=None, **kwargs):
        """
                        Инициализирует объект FirstFrameChoise.

                        Args:
                            *args: Дополнительные аргументы.
                            **kwargs: Дополнительные именованные аргументы.
                        """
        super().__init__(*args, **kwargs)
        self.db_manager = db_manager
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="Выберите филиал").grid(row=0, column=0, padx=10, pady=10,
                                                                         sticky='nsew')

        customtkinter.CTkLabel(master=self, text="Выберите структурное подразделение").grid(row=1, column=0, padx=10,
                                                                                            pady=10, sticky='nsew')
        data = self.db_manager.get_data("branch_office", "name", "")
        data = [str(row[0]) for row in data]
        self.combobox1_branch_office = customtkinter.CTkComboBox(master=self, values=[" "], state="readonly",
                                                                 command=self.load_data)
        self.combobox1_branch_office.grid(row=0, column=1, padx=10, pady=10, sticky="w")
        CTkToolTip(self.combobox1_branch_office, message="Выберите филиал.")
        self.FillComboBox(self.combobox1_branch_office, data)

        self.combobox2_structural_unit = customtkinter.CTkComboBox(master=self, values=[" "], state="readonly")
        self.combobox2_structural_unit.grid(row=1, column=1, padx=10, pady=10, sticky="w")
        CTkToolTip(self.combobox2_structural_unit, message="Выберите структурное подразделение.")

    def load_data(self, choice):
        """
                Загружает данные для структурного подразделения на основе выбранного филиала.

                Args:
                    choice (str): Выбранный филиал.
                """
        data = self.db_manager.exec_procedure("GetStructuralUnits", choice)
        data = [str(row[0]) for row in data]
        self.FillComboBox(self.combobox2_structural_unit, data)

    def FillComboBox(self, combobox, data_):
        """
                Заполняет комбобокс данными.

                Args:
                    combobox (customtkinter.CTkComboBox): Комбобокс, который нужно заполнить.
                    data_ (list): Список данных для заполнения.
                """
        data__ = [str(data) for data in data_]
        combobox.configure(values=data__)
        self.update()


class FirstFrameReport(customtkinter.CTkFrame):
    def __init__(self, master, first_frame_ch_instance, db_manager=None):
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="Отчет по статусам работы", fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10, sticky="ew")
        self.date_Start = CTkSelectDate.SelectDate(self, label_text="Выберите дату начала периода")
        self.date_Start.grid(row=1, columnspan=2, column=0, padx=10, pady=10, sticky="ew")

        self.date_End = CTkSelectDate.SelectDate(self, label_text="Выберите дату конца периода")
        self.date_End.grid(row=2, columnspan=2, column=0, padx=10, pady=10, sticky="ew")

        customtkinter.CTkLabel(master=self, text="Дата конца не включает последний день!").grid(row=3, columnspan=2,
                                                                                                column=0, padx=2,
                                                                                                pady=2, sticky="ew")

        self.MakeReport1Button = customtkinter.CTkButton(master=self, text="Сформировать отчет",
                                                         command=lambda: self.make_report())
        self.MakeReport1Button.grid(row=4, column=0, columnspan=2, pady=10, padx=10, sticky="ew")

    def compare_dates(self, date1_str, date2_str):
        return datetime.datetime.strptime(date1_str, "%Y-%m-%d") <= datetime.datetime.strptime(date2_str, "%Y-%m-%d")

    def make_report(self):
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()

        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return

        date_start_str = self.date_Start.get_current_date()
        date_end_str = self.date_End.get_current_date()

        if not date_start_str.strip():
            CTkMessagebox(title="Ошибка", message="Выберите дату начала периода!", icon="warning")
            return
        if not date_end_str.strip():
            CTkMessagebox(title="Ошибка", message="Выберите дату конца периода!", icon="warning")
            return

        try:
            date_start_iso = datetime.datetime.strptime(date_start_str, "%d.%m.%Y").strftime("%Y-%m-%d")
            date_end_iso = datetime.datetime.strptime(date_end_str, "%d.%m.%Y").strftime("%Y-%m-%d")
        except ValueError as e:
            CTkMessagebox(title="Ошибка", message="Неверный формат даты!", icon="warning")
            return



        try:
            date_start = datetime.datetime.strptime(date_start_iso, "%Y-%m-%d")
            date_end = datetime.datetime.strptime(date_end_iso, "%Y-%m-%d")
        except ValueError as e:
            CTkMessagebox(title="Ошибка", message="Неверный формат даты!", icon="warning")
            return

        if self.compare_dates(date_start_iso, date_end_iso):
            diff = date_end - date_start
            if diff.days >= 0:
                print(branch_office)
                print(structural_unit)
                self.make_document(date_start, date_end)
            else:
                CTkMessagebox(title="Ошибка", message="Дата начала позже даты окончания.", icon="warning")
        else:
            CTkMessagebox(title="Ошибка", message="Дата начала позже даты окончания.", icon="warning")

    def make_document(self, date_start, date_end):
        data = self.db_manager.exec_procedure("GetBasicInfoStatusBetweenDatesInBranchAndStructUnit",
                                              date_start, date_end,
                                              f"{self.first_frame_choise.combobox1_branch_office.get()}",
                                              f"{self.first_frame_choise.combobox2_structural_unit.get()}")
        root = tk.Tk()
        root.withdraw()
        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word Document", "*.docx")],
                                                 title="Выберите место сохранения документа",
                                                 initialfile="ОтчетПоСтатусамРаботы")

        if not file_path:
            return

        try:
            doc = Document()
            title = doc.add_heading(
                f'Отчет по статусам работы устройств.\n Филиал - {self.first_frame_choise.combobox1_branch_office.get()} \n '
                f'Структурное подразделение - {self.first_frame_choise.combobox2_structural_unit.get()}', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Название'
            hdr_cells[1].text = 'Сетевое имя'
            hdr_cells[2].text = 'Статус'
            hdr_cells[3].text = 'Место установки'
            hdr_cells[4].text = 'Дата'

            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = item[0]
                row_cells[1].text = item[1]
                row_cells[2].text = "On" if item[2] else "Off"
                row_cells[3].text = item[4]
                row_cells[4].text = item[3].strftime("%d.%m.%Y %H:%M:%S")

                color = RGBColor(0, 128, 0) if item[2] else RGBColor(255, 0, 0)
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = color

            doc.save(file_path)
            CTkMessagebox(title="Успех", message="Отчет успешно сформирован.", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message="Ошибка при формировании отчета!", icon="warning")

class SecondFrameReport(customtkinter.CTkFrame):
    """
    Класс для создания фрейма формирования отчета по ремонтам.

    Attributes:
        first_frame_choise (FirstFrameChoise): Экземпляр класса FirstFrameChoise для доступа к выбранным данным.
        second_frame_instance: Экземпляр класса SecondFrame.

    Methods:
        compare_dates: Сравнивает две даты на предмет их порядка.
        make_report: Формирует отчет по ремонтам на основе выбранных параметров.
        make_document: Создает документ с отчетом по ремонтам на основе полученных данных.
    """

    def __init__(self, master, first_frame_ch_instance, second_instance, db_manager=None):
        """
        Инициализирует объект SecondFrameReport.

        Args:
            master: Родительский виджет.
            first_frame_ch_instance (FirstFrameChoise): Экземпляр класса FirstFrameChoise.
            second_instance: Экземпляр класса SecondFrame.
        """
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.second_frame_instance = second_instance
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="Отчет по ремонтам", fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10, sticky="ew")
        self.date_Start = CTkSelectDate.SelectDate(self, label_text="Выберите дату начала периода")
        self.date_Start.grid(row=1, columnspan=2, column=0, padx=10, pady=10, sticky="ew")

        self.date_End = CTkSelectDate.SelectDate(self, label_text="Выберите дату конца периода")
        self.date_End.grid(row=2, columnspan=2, column=0, padx=10, pady=10, sticky="ew")

        customtkinter.CTkLabel(master=self, text="Дата конца не включает последний день!").grid(row=3, columnspan=2,
                                                                                                column=0, padx=2,
                                                                                                pady=2, sticky="ew")

        self.MakeReport1Button = customtkinter.CTkButton(master=self, text="Сформировать отчет",
                                                         command=lambda: self.make_report())
        self.MakeReport1Button.grid(row=4, column=0, columnspan=2, pady=10, padx=10, sticky="ew")

    def make_report(self):
        """
        Сравнивает две даты на предмет их порядка.
        """
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()

        # Check if either combobox is empty
        if not branch_office.strip():
            CTkMessagebox.showwarning("Ошибка", "Выберите филиал!")
            return
        if not structural_unit.strip():
            CTkMessagebox.showwarning("Ошибка", "Выберите структурное подразделение!")
            return

        date_start_str = self.date_Start.get_current_date()
        date_end_str = self.date_End.get_current_date()

        # Check if either date is empty
        if not date_start_str.strip():
            CTkMessagebox.showwarning("Ошибка", "Выберите дату начала периода!")
            return
        if not date_end_str.strip():
            CTkMessagebox.showwarning("Ошибка", "Выберите дату конца периода!")
            return

        try:
            # Преобразование даты из дд.мм.гггг в гггг-мм-дд
            date_start_iso = datetime.datetime.strptime(date_start_str, "%d.%m.%Y").strftime("%Y-%m-%d")
            date_end_iso = datetime.datetime.strptime(date_end_str, "%d.%m.%Y").strftime("%Y-%m-%d")
        except ValueError as e:
            CTkMessagebox.showwarning("Ошибка", "Неверный формат даты!")
            return

        try:
            date_start = datetime.datetime.strptime(date_start_iso, "%Y-%m-%d")
            date_end = datetime.datetime.strptime(date_end_iso, "%Y-%m-%d")
        except ValueError as e:
            CTkMessagebox.showwarning("Ошибка", "Неверный формат даты!")
            return

        if self.compare_dates(date_start_iso, date_end_iso):
            diff = date_end - date_start
            if diff.days >= 0:
                print(branch_office)
                print(structural_unit)
                self.make_document()
            else:
                CTkMessagebox.showwarning("Ошибка", "Дата начала позже даты окончания.")
        else:
            CTkMessagebox.showwarning("Ошибка", "Дата начала позже даты окончания.")

    def compare_dates(self, date1_str, date2_str):
        return datetime.datetime.strptime(date1_str, "%Y-%m-%d") <= datetime.datetime.strptime(date2_str, "%Y-%m-%d")

    def make_document(self):
        """Формирует отчет по ремонтам на основе выбранных параметров."""
        try:
            date_start_str = self.date_Start.get_current_date()
            date_end_str = self.date_End.get_current_date()

            date_start = datetime.datetime.strptime(date_start_str, "%d.%m.%Y").strftime("%Y-%m-%d")
            date_end = datetime.datetime.strptime(date_end_str, "%d.%m.%Y").strftime("%Y-%m-%d")

            data = self.db_manager.exec_procedure("GetRepairsBetweenTwoDates", date_start, date_end,
                                             f"{self.first_frame_choise.combobox1_branch_office.get()}",
                                             f"{self.first_frame_choise.combobox2_structural_unit.get()}")
            root = tk.Tk()
            root.withdraw()  # Скрытие корневого окна
            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word Document", "*.docx")],
                                                     title="Выберите место сохранения документа",
                                                     initialfile="ОтчетПоРемонтам")
            if not file_path:
                return  # Прерывание функции, если пользователь не выбрал место сохранения
            doc = Document()
            title = doc.add_heading(
                f'Отчет по ремонтам устройств.\n Филиал - {self.first_frame_choise.combobox1_branch_office.get()} \n '
                f'Структурное подразделение - {self.first_frame_choise.combobox2_structural_unit.get()}', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Добавление таблицы с данными
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Заголовки столбцов
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Сетевое имя'
            hdr_cells[1].text = 'IP'
            hdr_cells[2].text = 'Описание'
            hdr_cells[3].text = 'Дата ремонта'
            hdr_cells[4].text = 'Папка с документами'

            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item[0])  # Название
                row_cells[1].text = str(item[1])  # Сетевое имя
                row_cells[2].text = str(item[2])
                row_cells[3].text = datetime.datetime.strptime(str(item[3]), "%Y-%m-%d").strftime("%d.%m.%Y")  # Место установки
                row_cells[4].text = str(item[4])

            # Сохранение документа
            doc.save(file_path)
            CTkMessagebox(title="Успех", message="Отчет успешно сформирован.", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка при формировании отчета!\n{e}", icon="warning")


class ThirdFrameReport(customtkinter.CTkFrame):
    """
        Класс для создания фрейма формирования отчета по выбранному устройству.

        Attributes:
            first_frame_choise (FirstFrameChoise): Экземпляр класса FirstFrameChoise для доступа к выбранным данным.
            second_frame_instance: Экземпляр класса SecondFrame.

        Methods:
            load_devices_types: Загружает типы устройств на основе выбранных данных.
            load_devices: Загружает устройства выбранного типа на основе выбранных данных.
            make_report: Формирует отчет по выбранному устройству.
        """

    def __init__(self, master, first_frame_ch_instance, second_instance,db_manager=None):
        """
                Инициализирует объект ThirdFrameReport.

                Args:
                    master: Родительский виджет.
                    first_frame_ch_instance (FirstFrameChoise): Экземпляр класса FirstFrameChoise.
                    second_instance: Экземпляр класса SecondFrame.
                """
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.second_frame_instance = second_instance
        self.grid_rowconfigure((1, 2, 3, 4), weight=1)
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="Отчет по выбранному устройству", fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10, sticky="ew")
        customtkinter.CTkButton(master=self, text="Загрузить устройства",
                                command=lambda: self.load_devices_types()).grid(
            row=1, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

        customtkinter.CTkLabel(master=self, text="Выберите тип устройства").grid(
            row=2, column=0, padx=10, pady=10)

        self.type_of_device_combob = customtkinter.CTkComboBox(master=self, values=[" "], state="readonly",
                                                               command=self.load_devices)
        self.type_of_device_combob.grid(row=2, column=1, padx=10, pady=10)

        customtkinter.CTkLabel(master=self, text="Выберите устройство").grid(
            row=3, column=0, padx=10, pady=10)
        self.device_combob = customtkinter.CTkComboBox(master=self, values=[" "], state="readonly")
        self.device_combob.grid(row=3, column=1, padx=10, pady=10)

        customtkinter.CTkButton(master=self, text="Сформировать отчёт",
                                command=lambda: self.make_report()).grid(
            row=4, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

    def load_devices_types(self):
        """Загружает типы устройств на основе выбранных данных."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()
        # Check if either combobox is empty
        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return
        try:
            id_branch = self.db_manager.get_data("branch_office", "id", f"name = '{branch_office}'")[0][0]
            id_unit = self.db_manager.get_data("structural_unit", "id", f"name = '{structural_unit}'")[0][0]
            data = self.db_manager.exec_procedure("GetUniqueDeviceNames", id_branch, id_unit)
            data = [str(row[0]) for row in data]
            self.type_of_device_combob.configure(values=data)
            self.update()
            CTkMessagebox(title="Успех", message="Типы устройств успешно загуржены", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка на этапе загрузки типов устройств!\n {e}", icon="warning")

    def load_devices(self, choice):
        """Загружает устройства выбранного типа на основе выбранных данных."""
        print(choice)
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()
        # Check if either combobox is empty
        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return
        id_branch = self.db_manager.get_data("branch_office", "id", f"name = '{branch_office}'")[0][0]
        id_unit = self.db_manager.get_data("structural_unit", "id", f"name = '{structural_unit}'")[0][0]
        id_type_of_device = self.db_manager.get_data("type_of_device", "id", f"name = '{choice}'")[0][0]

        data = self.db_manager.exec_procedure("GetDeviceDetails", id_branch, id_unit, id_type_of_device)
        data = [str(row[0] + " | " + row[1]) for row in data]
        self.device_combob.configure(values=data)
        self.update()

    def make_report(self):
        """Формирует документ отчета по выбранному устройству."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()
        type_of_device = self.type_of_device_combob.get()
        device = self.device_combob.get()

        # Check if either combobox is empty
        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return
        if not type_of_device.strip():
            CTkMessagebox(title="Ошибка", message="Выберите тип устройства!", icon="warning")
            return
        if not device.strip():
            CTkMessagebox(title="Ошибка", message="Выберите устройствo!", icon="warning")
            return
        parts = device.split('|')
        # Удаление лишних пробелов с каждой части
        ip_address = parts[0].strip()
        network_name = parts[1].strip()
        try:
            root = tk.Tk()
            root.withdraw()  # Скрытие корневого окна
            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word Document", "*.docx")],
                                                     title="Выберите место сохранения документа",
                                                     initialfile=f"ОтчетПоУстройству_{network_name}")

            if not file_path:
                return  # Прерывание функции, если пользователь не выбрал место сохранения

            data = self.db_manager.exec_procedure("GetDeviceInfoByBranchStrUnitTypeOfDeviceIpNetworkName",
                                             branch_office, structural_unit, type_of_device, ip_address, network_name)
            print(data)
            doc = DocxTemplate("resources\\pattern_for_third_report.docx")
            context = {
                'branch_office': data[0][0],
                'structural_unit': data[0][1],
                'inv_numb': data[0][2],
                'name': data[0][3],
                'network_name': data[0][4],
                'type_of_device': data[0][5],
                'location': data[0][6],
                'description_basic': data[0][7],
                'material_resp_person': data[0][8],
                'last_status': 'Вкл' if data[0][9] else 'Выкл',
                'last_repair': data[0][10] if data[0][10] is not None else '',
                'serial_numb': data[0][11],
                'mac_adr': data[0][12],
                'os': data[0][13],
                'year_of_pushare': data[0][14],
                'month_of_garanty': data[0][15],
                'processor': data[0][16],
                'ram': data[0][17],
                'motherboard': data[0][18],
                'graphicCard': data[0][19],
                'psu': data[0][20],
                'networkCard': data[0][21],
                'cooler': data[0][22],
                'chasis': data[0][23],
                'hdd': data[0][24],
                'ssd': data[0][25],
                'monitor': data[0][26],
                'keyboard': data[0][27],
                'mouse': data[0][28],
                'audio': data[0][29]
            }
            doc.render(context)
            doc.save(file_path)
            CTkMessagebox(title="Успех", message="Отчет успешно сформирован.", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message="Ошибка при формировании отчета!\n{e}", icon="warning")


class FourthFrameReport(customtkinter.CTkFrame):
    """Класс для создания фрейма формирования отчета по месту установки.

        Attributes:
            first_frame_choise (FirstFrameChoise): Экземпляр класса FirstFrameChoise для доступа к выбранным данным.

        Methods:
            load_locations: Загружает доступные места установки на основе выбранных данных.
            make_report: Формирует отчет по выбранному месту установки.
        """

    def __init__(self, master, first_frame_ch_instance,db_manager=None):
        """
                Инициализирует объект FourthFrameReport.

                Args:
                    master: Родительский виджет.
                    first_frame_ch_instance (FirstFrameChoise): Экземпляр класса FirstFrameChoise.
                """
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.grid_rowconfigure((1, 2, 3, 4), weight=1)
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="Отчет по месту установки", fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10, sticky="ew")
        customtkinter.CTkButton(master=self, text="Загрузить места установки",
                                command=lambda: self.load_locations()).grid(
            row=1, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

        customtkinter.CTkLabel(master=self, text="Выберите место установки").grid(
            row=2, column=0, padx=10, pady=10)

        self.place_of_install_combob = customtkinter.CTkComboBox(master=self, values=[" "], state="readonly")
        self.place_of_install_combob.grid(row=2, column=1, padx=10, pady=10)

        customtkinter.CTkButton(master=self, text="Сформировать отчет",
                                command=lambda: self.make_report()).grid(
            row=3, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

    def load_locations(self):
        """Загружает доступные места установки на основе выбранных данных."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()
        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return
        try:

            data = self.db_manager.exec_procedure("GetPlaceOfInstallationByBranchAndUnit", branch_office, structural_unit)
            data = [str(row[0]) for row in data]
            self.place_of_install_combob.configure(values=data)
            self.update()
            CTkMessagebox(title="Успех", message="Типы устройств успешно загуржены", icon="check")

        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка на этапе загрузки типов устройств!\n {e}", icon="warning")

    def make_report(self):
        """Формирует отчет по выбранному месту установки."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()
        place_of_install = self.place_of_install_combob.get()

        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return
        if not place_of_install.strip():
            CTkMessagebox(title="Ошибка", message="Выберите место установки!", icon="warning")
            return

        try:
            data = self.db_manager.exec_procedure("GetDeviceInfoByBranchStrUnitPlaceOf", branch_office, structural_unit,
                                             place_of_install)
            print(data)

            wb = Workbook()
            ws = wb.active

            # Добавляем строки с данными о филиале, структурном подразделении и месте установки
            ws.append(["Филиал:", branch_office])
            ws.append(["Структурное подразделение:", structural_unit])
            ws.append(["Место установки:", place_of_install])
            ws.append([])  # Пустая строка для разделения

            # Заголовки столбцов
            headers = [
                'Филиал',
                'Структурное подразделение',
                'Инвентарный номер',
                'Наименование устройства',
                'Сетевое имя',
                'Тип устройства',
                'Установлен в',
                'Описание',
                'Материально ответственное лицо',
                'Последний статус',
                'Дата последнего ремонта',
                'Серийный номер',
                'MAC адрес',
                'Операционная система',
                'Год покупки',
                'Месяцы гарантии',
                'Процессор',
                'ОЗУ',
                'Материнская плата',
                'Видеокарта',
                'Блок питания',
                'Сетевая карта',
                'Кулер',
                'Корпус',
                'HDD',
                'SSD',
                'Монитор',
                'Клавиатура',
                'Мышь',
                'Звуковые устройства'
            ]

            # Добавляем заголовки в пятую строку (т.к. первые четыре строки заняты информацией о филиале)
            for col, header in enumerate(headers, start=1):
                ws.cell(row=5, column=col, value=header)

            # Автоподбор ширины столбцов на основе заголовков и данных
            column_widths = [len(header) for header in headers]  # Инициализируем ширины столбцов на основе заголовков

            # Заполнение данными начиная с шестой строки
            for row_num, item in enumerate(data, start=6):
                for col_num, value in enumerate(item, start=1):
                    ws.cell(row=row_num, column=col_num, value=value)
                    column_widths[col_num - 1] = max(column_widths[col_num - 1], len(str(value)))

            # Настройка ширины столбцов
            for col_num, width in enumerate(column_widths, start=1):
                column_letter = get_column_letter(col_num)
                ws.column_dimensions[column_letter].width = (
                                                                    width + 2) * 1.2  # Увеличиваем ширину на 20% от максимальной длины

            root = tk.Tk()
            root.withdraw()  # Скрыть корневое окно

            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                     filetypes=[("Excel files", "*.xlsx")],
                                                     title="Выберите место сохранения Excel-файла",
                                                     initialfile=f"Отчет_по_устройствам_{branch_office}_{structural_unit}_{place_of_install}.xlsx")

            if file_path:
                wb.save(file_path)
                CTkMessagebox(title="Успех", message="Отчет успешно сохранен!", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка на этапе формирования отчета!\n {e}", icon="warning")


class FifthFrameReport(customtkinter.CTkFrame):
    """Класс для создания фрейма формирования отчета по устройствам филиала.

        Attributes:
            first_frame_choise (FirstFrameChoise): Экземпляр класса FirstFrameChoise для доступа к выбранным данным.

        Methods:
            make_report: Формирует отчет по устройствам выбранного филиала.
        """

    def __init__(self, master, first_frame_ch_instance,db_manager=None):
        """Инициализирует объект FifthFrameReport.

                    Args:
                        master: Родительский виджет.
                        first_frame_ch_instance (FirstFrameChoise): Экземпляр класса FirstFrameChoise.
                    """
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.grid_rowconfigure((1, 2, 3, 4), weight=1)
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="    Отчет по устройствам филиала    ", fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10,
                                                        sticky="ew")
        customtkinter.CTkButton(master=self, text="Сформировать отчет",
                                command=lambda: self.make_report()).grid(
            row=1, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

    def make_report(self):
        """Формирует отчет по устройствам выбранного филиала."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()

        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return

        try:
            data = self.db_manager.exec_procedure("GetAllDevicesByBranch", branch_office)
            print(data)

            wb = Workbook()
            ws = wb.active

            # Добавляем строку с данными о филиале
            ws.append(["Филиал:", branch_office])
            ws.append([])  # Пустая строка для разделения

            # Заголовки столбцов
            headers = [
                'Филиал',
                'Структурное подразделение',
                'Инвентарный номер',
                'Наименование устройства',
                'Сетевое имя',
                'Тип устройства',
                'Установлен в',
                'Описание',
                'Материально ответственное лицо',
                'Последний статус',
                'Дата последнего ремонта',
                'Серийный номер',
                'MAC адрес',
                'Операционная система',
                'Год покупки',
                'Месяцы гарантии',
                'Процессор',
                'ОЗУ',
                'Материнская плата',
                'Видеокарта',
                'Блок питания',
                'Сетевая карта',
                'Кулер',
                'Корпус',
                'HDD',
                'SSD',
                'Монитор',
                'Клавиатура',
                'Мышь',
                'Звуковые устройства'
            ]

            # Добавляем заголовки в третью строку (т.к. первые две строки заняты филиалом и пустой строкой)
            for col, header in enumerate(headers, start=1):
                ws.cell(row=3, column=col, value=header)

            # Автоподбор ширины столбцов на основе заголовков и данных
            column_widths = [len(header) for header in
                             headers]  # Инициализируем ширины столбцов на основе заголовков

            # Заполнение данными начиная с четвертой строки
            for row_num, item in enumerate(data, start=4):
                for col_num, value in enumerate(item, start=1):
                    ws.cell(row=row_num, column=col_num, value=value)
                    column_widths[col_num - 1] = max(column_widths[col_num - 1], len(str(value)))

            # Настройка ширины столбцов
            for col_num, width in enumerate(column_widths, start=1):
                column_letter = get_column_letter(col_num)
                ws.column_dimensions[column_letter].width = (
                                                                    width + 2) * 1.2  # Увеличиваем ширину на 20% от максимальной длины

            root = tk.Tk()
            root.withdraw()  # Скрыть корневое окно

            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                     filetypes=[("Excel files", "*.xlsx")],
                                                     title="Выберите место сохранения Excel-файла",
                                                     initialfile=f"Отчет_по_устройствам_{branch_office}.xlsx")

            if file_path:
                wb.save(file_path)
                CTkMessagebox(title="Успех", message="Отчет успешно сохранен!", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка на этапе формирования отчета!\n {e}", icon="warning")


class SixthFrameReport(customtkinter.CTkFrame):
    """Класс для создания фрейма формирования отчета по устройствам структурного подразделения.

        Attributes:
            first_frame_choise (FirstFrameChoise): Экземпляр класса FirstFrameChoise для доступа к выбранным данным.

        Methods:
            make_report: Формирует отчет по устройствам выбранного структурного подразделения.
        """

    def __init__(self, master, first_frame_ch_instance,db_manager=None):
        """Инициализирует объект SixthFrameReport.

                    Args:
                        master: Родительский виджет.
                        first_frame_ch_instance (FirstFrameChoise): Экземпляр класса FirstFrameChoise.
                    """
        super().__init__(master)
        self.db_manager = db_manager
        self.first_frame_choise = first_frame_ch_instance
        self.grid_rowconfigure((1, 2, 3, 4), weight=1)
        self.configure(border_color="dodgerblue", border_width=3)
        customtkinter.CTkLabel(master=self, text="    Отчет по устройствам структурного подразделения    ",
                               fg_color="gray30",
                               font=("Arial", 14)).grid(row=0, columnspan=3, column=0, padx=10, pady=10,
                                                        sticky="ew")
        customtkinter.CTkButton(master=self, text="Сформировать отчет",
                                command=lambda: self.make_report()).grid(
            row=1, column=0, padx=10, pady=10, sticky="ew", columnspan=3)

    def make_report(self):
        """Формирует отчет по устройствам выбранного структурного подразделения."""
        branch_office = self.first_frame_choise.combobox1_branch_office.get()
        structural_unit = self.first_frame_choise.combobox2_structural_unit.get()

        if not branch_office.strip():
            CTkMessagebox(title="Ошибка", message="Выберите филиал!", icon="warning")
            return
        if not structural_unit.strip():
            CTkMessagebox(title="Ошибка", message="Выберите структурное подразделение!", icon="warning")
            return

        try:
            data = self.db_manager.exec_procedure("GetAllDevicesByUnit", branch_office, structural_unit)
            print(data)

            wb = Workbook()
            ws = wb.active

            # Добавляем строки с данными о филиале и структурном подразделении
            ws.append(["Филиал:", branch_office])
            ws.append(["Структурное подразделение:", structural_unit])
            ws.append([])  # Пустая строка для разделения

            # Заголовки столбцов
            headers = [
                'Филиал',
                'Структурное подразделение',
                'Инвентарный номер',
                'Наименование устройства',
                'Сетевое имя',
                'Тип устройства',
                'Установлен в',
                'Описание',
                'Материально ответственное лицо',
                'Последний статус',
                'Дата последнего ремонта',
                'Серийный номер',
                'MAC адрес',
                'Операционная система',
                'Год покупки',
                'Месяцы гарантии',
                'Процессор',
                'ОЗУ',
                'Материнская плата',
                'Видеокарта',
                'Блок питания',
                'Сетевая карта',
                'Кулер',
                'Корпус',
                'HDD',
                'SSD',
                'Монитор',
                'Клавиатура',
                'Мышь',
                'Звуковые устройства'
            ]

            # Добавляем заголовки в пятую строку (т.к. первые четыре строки заняты информацией о филиале и подразделении)
            for col, header in enumerate(headers, start=1):
                ws.cell(row=5, column=col, value=header)

            # Автоподбор ширины столбцов на основе заголовков и данных
            column_widths = [len(header) for header in
                             headers]  # Инициализируем ширины столбцов на основе заголовков

            # Заполнение данными начиная с шестой строки
            for row_num, item in enumerate(data, start=6):
                for col_num, value in enumerate(item, start=1):
                    ws.cell(row=row_num, column=col_num, value=value)
                    column_widths[col_num - 1] = max(column_widths[col_num - 1], len(str(value)))

            # Настройка ширины столбцов
            for col_num, width in enumerate(column_widths, start=1):
                column_letter = get_column_letter(col_num)
                ws.column_dimensions[column_letter].width = (
                                                                    width + 2) * 1.2  # Увеличиваем ширину на 20% от максимальной длины

            root = tk.Tk()
            root.withdraw()  # Скрыть корневое окно

            file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                     filetypes=[("Excel files", "*.xlsx")],
                                                     title="Выберите место сохранения Excel-файла",
                                                     initialfile=f"Отчет_по_устройствам_{branch_office}_{structural_unit}.xlsx")

            if file_path:
                wb.save(file_path)
                CTkMessagebox(title="Успех", message="Отчет успешно сохранен!", icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", message=f"Ошибка на этапе формирования отчета!\n {e}", icon="warning")


if __name__ == "__main__":
    root = tkinter.Tk()
    app = Reports(root)
    root.mainloop()
